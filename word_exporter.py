"""
Word 文档导出模块 - 生成包含凭证图片的报销明细 Word 文档
"""
import os
from datetime import datetime
from io import BytesIO

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from logger import get_logger

log = get_logger("word_exporter")

UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")


def _load_image_for_docx(full_path: str) -> BytesIO | None:
    """用 Pillow 读取图片并统一转为 PNG 流，兼容各种上传格式

    Args:
        full_path: 图片文件绝对路径

    Returns:
        PNG 格式的 BytesIO 流，失败返回 None
    """
    try:
        img = Image.open(full_path)
        # 统一转为 RGB（处理 RGBA、调色板等模式）
        if img.mode in ("RGBA", "LA", "P"):
            # 保留透明通道用 RGBA
            if img.mode == "P":
                img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as e:
        log.warning("Pillow 读取图片失败 %s: %s", full_path, e)
        return None


def export_to_word(expenses: list[dict]) -> str:
    """生成 Word 文档，包含报销明细 + 凭证图片 + 总计

    Args:
        expenses: 报销记录列表，每条包含
            date, person, category, amount, remark, image_path, username

    Returns:
        生成的 Word 文件路径
    """
    doc = Document()

    # ---- 标题 ----
    title = doc.add_heading("报销明细汇总表", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)

    # ---- 汇总信息 ----
    total_amount = round(sum(e.get("amount", 0) for e in expenses), 2)
    total_count = len(expenses)

    summary = doc.add_paragraph()
    summary.add_run(f"共 {total_count} 条记录，合计金额：¥{total_amount:.2f}").bold = True

    doc.add_paragraph("")  # 空行

    # ---- 按报销人分组 ----
    person_map: dict[str, list] = {}
    for exp in expenses:
        person = exp.get("person", "未知")
        if person not in person_map:
            person_map[person] = []
        person_map[person].append(exp)

    # ---- 逐人生成明细 ----
    for person, records in person_map.items():
        person_total = round(sum(r.get("amount", 0) for r in records), 2)

        # 报销人标题
        doc.add_heading(f"报销人：{person}（{len(records)} 条，¥{person_total:.2f}）", level=1)

        # 明细表格
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        hdr = table.rows[0].cells
        headers = ["日期", "品类", "金额", "备注", "凭证图片"]
        for i, h in enumerate(headers):
            p = hdr[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)

        # 数据行
        for r in records:
            row = table.add_row().cells
            row[0].text = str(r.get("date", ""))
            row[1].text = str(r.get("category", ""))
            row[2].text = f"¥{r.get('amount', 0):.2f}"
            row[3].text = str(r.get("remark", "") or "")

            # 插入凭证图片
            img_cell = row[4]
            img_path = r.get("image_path", "")
            if img_path:
                full_path = os.path.join(UPLOAD_FOLDER, os.path.basename(img_path))
                if os.path.isfile(full_path):
                    # 用 Pillow 统一转 PNG 再嵌入，兼容各种上传格式
                    img_stream = _load_image_for_docx(full_path)
                    if img_stream:
                        try:
                            p = img_cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(img_stream, width=Inches(2.0))
                        except Exception as e:
                            log.warning("插入图片失败 %s: %s", full_path, e)
                            img_cell.text = "图片加载失败"
                    else:
                        img_cell.text = "图片格式不支持"
                else:
                    img_cell.text = "图片已过期"
            else:
                img_cell.text = "-"

            # 设置单元格字体大小
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        # 小计
        sub_total = doc.add_paragraph()
        sub_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sub_total.add_run(f"{person} 小计：¥{person_total:.2f}").bold = True

        doc.add_page_break()

    # ---- 总计页 ----
    doc.add_heading("总计", level=1)
    total_p = doc.add_paragraph()
    total_p.add_run(f"报销总笔数：{total_count} 条\n").font.size = Pt(12)
    total_p.add_run(f"报销总金额：¥{total_amount:.2f}").font.size = Pt(14)

    # 各人小计汇总表
    doc.add_paragraph("")
    doc.add_heading("各人汇总", level=2)

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    for i, h in enumerate(["报销人", "笔数", "金额"]):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True

    for person, records in sorted(person_map.items()):
        row = summary_table.add_row().cells
        row[0].text = person
        row[1].text = str(len(records))
        row[2].text = f"¥{sum(r.get('amount', 0) for r in records):.2f}"

    # 保存文件
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "exports")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(
        output_dir,
        f"报销明细_含凭证_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    doc.save(output_path)
    log.info("Word 文档已生成: %s", output_path)
    return output_path
